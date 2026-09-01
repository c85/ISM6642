#!/usr/bin/env python3
"""Update the existing assignment reports from the current analysis outputs."""

from __future__ import annotations

import csv
import json
import zipfile
from collections import Counter
from copy import deepcopy
from pathlib import Path

import pandas as pd
from docx import Document


ROOT = Path(__file__).resolve().parent.parent
DATA_DIR = ROOT / "data"
OUTPUT_DIR = ROOT / "outputs"
FIGURE_DIR = OUTPUT_DIR / "figures"
REPORT_DIR = ROOT / "reports"


def matches_prefix(text: str, prefix: str | tuple[str, ...]) -> bool:
    prefixes = (prefix,) if isinstance(prefix, str) else prefix
    return any(text.strip().startswith(candidate) for candidate in prefixes)


def set_paragraph(doc: Document, prefix: str | tuple[str, ...], text: str) -> None:
    for paragraph in doc.paragraphs:
        if matches_prefix(paragraph.text, prefix):
            paragraph.text = text
            return
    raise ValueError(f"Could not find paragraph beginning with {prefix!r}")


def set_paragraph_after(
    doc: Document,
    heading_prefix: str,
    prefix: str | tuple[str, ...],
    text: str,
) -> None:
    heading_found = False
    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith(heading_prefix):
            heading_found = True
            continue
        if heading_found and matches_prefix(paragraph.text, prefix):
            paragraph.text = text
            return
    raise ValueError(
        f"Could not find paragraph beginning with {prefix!r} after {heading_prefix!r}"
    )


def set_cell(cell, value: object) -> None:
    cell.text = str(value)


def set_table(table, rows: list[list[object]]) -> None:
    if any(len(row) != len(table.columns) for row in rows):
        raise ValueError("Replacement table column count does not match the existing report table")
    while len(table.rows) > len(rows):
        table._tbl.remove(table.rows[-1]._tr)
    while len(table.rows) < len(rows):
        table._tbl.append(deepcopy(table.rows[-1]._tr))
    for table_row, values in zip(table.rows, rows):
        for cell, value in zip(table_row.cells, values):
            set_cell(cell, value)


def money(value: float, decimals: int = 2) -> str:
    return f"${value:,.{decimals}f}"


def p_value(value: float) -> str:
    return f"{value:.3f}" if value >= 0.001 else f"{value:.5f}"


def replace_media(docx_path: Path, image_paths: list[Path]) -> None:
    """Replace embedded report images while retaining their existing placement."""
    temporary = docx_path.with_suffix(".updated.docx")
    with zipfile.ZipFile(docx_path, "r") as source, zipfile.ZipFile(
        temporary, "w", compression=zipfile.ZIP_DEFLATED
    ) as destination:
        replacements = {
            f"word/media/image{index}.png": image.read_bytes()
            for index, image in enumerate(image_paths, start=1)
        }
        for item in source.infolist():
            data = replacements.get(item.filename, source.read(item.filename))
            destination.writestr(item, data)
    temporary.replace(docx_path)


def update_eda_report() -> None:
    path = REPORT_DIR / "Florida_Fitness_Market_EDA_Report.docx"
    doc = Document(path)
    location_counts = Counter(
        row["GymChain"]
        for row in csv.DictReader((DATA_DIR / "gym_locations.csv").open(encoding="utf-8"))
    )
    summary = pd.read_csv(OUTPUT_DIR / "summary_statistics.csv", index_col=0)
    correlations = pd.read_csv(OUTPUT_DIR / "correlation_matrix.csv", index_col=0)[
        "MembershipPrice"
    ].drop("MembershipPrice")
    outliers = pd.read_csv(OUTPUT_DIR / "outliers.csv")

    set_paragraph(
        doc,
        "All 67 Florida counties",
        "All 67 Florida counties | Census + Google Places + official gym pricing | OLS regression",
    )
    set_paragraph(
        doc,
        "Prepared for the MSIS",
        "Prepared for the MSIS Data Science Assignment | September 1, 2026",
    )
    set_paragraph(
        doc,
        (
            "Observed demographic and location data",
            "Observed demographics, locations, and official advertised price observations",
        ),
        "Observed demographics, locations, and official advertised price observations; transparent estimates only where no local offer was available",
    )
    set_paragraph(
        doc,
        "This report examines",
        "This report examines the market structure of all 67 Florida counties before regression modeling. It combines official demographic measures, live Google Places results for LA Fitness and four competitor brands, and official plan-level advertised prices where public enrollment pages expose them. The objective is to identify distributions, relationships, and unusual counties that may affect model interpretation or expansion decisions.",
    )
    set_paragraph(
        doc,
        "The Google Places Text Search",
        "The Google Places Text Search pipeline issued branded, county-specific queries, filtered results to exact county and brand matches, and deduplicated by Place ID. The refreshed location file retains each gym's Google display name and formatted street address, along with coordinates, rating, and review count. It contains 428 locations across LA Fitness, Planet Fitness, YouFit, Gold's Gym, and Orangetheory Fitness.",
    )
    set_paragraph(
        doc,
        "Google Places results",
        "Google Places results are a retrieval-time snapshot. Search relevance, closures, brand naming, and data coverage can change, so counts should be validated before a real estate commitment. Official pricing is also a retrieval-time snapshot of advertised offers, not a transaction-price survey.",
    )
    set_paragraph(
        doc,
        "Figure 1.",
        "Figure 1. County population versus advertised/estimated monthly membership-price target.",
    )
    set_paragraph(
        doc,
        "Figure 2.",
        "Figure 2. Median household income versus advertised/estimated monthly membership-price target.",
    )
    set_paragraph(
        doc,
        "Figure 3.",
        "Figure 3. Observed LA Fitness location count versus the county price target.",
    )
    set_paragraph(
        doc,
        "Figure 4.",
        "Figure 4. Distribution of county membership-price targets.",
    )

    set_paragraph(
        doc,
        "CountyPopulation:",
        "CountyPopulation: 2020 Census PL 94-171 total population (P1_001N).",
    )
    set_paragraph(
        doc,
        "MedianHouseholdIncome:",
        "MedianHouseholdIncome: 2021 ACS five-year estimate (B19013_001E).",
    )
    set_paragraph_after(doc, "Correlation Findings", "PopulationDensity:", f"PopulationDensity: r = {correlations['PopulationDensity']:.3f} (positive).")
    set_paragraph_after(doc, "Correlation Findings", "LAFitnessLocations:", f"LAFitnessLocations: r = {correlations['LAFitnessLocations']:.3f} (negative).")
    set_paragraph_after(doc, "Correlation Findings", "CompetitorLocations:", f"CompetitorLocations: r = {correlations['CompetitorLocations']:.3f} (positive).")
    set_paragraph_after(doc, "Correlation Findings", "CountyPopulation:", f"CountyPopulation: r = {correlations['CountyPopulation']:.3f} (positive).")
    set_paragraph_after(doc, "Correlation Findings", "MedianHouseholdIncome:", f"MedianHouseholdIncome: r = {correlations['MedianHouseholdIncome']:.3f} (positive and strongest).")
    set_paragraph_after(doc, "Correlation Findings", "AvgGymRating:", f"AvgGymRating: r = {correlations['AvgGymRating']:.3f} (positive).")
    set_paragraph(
        doc,
        "Influence diagnostics flagged",
        "Influence diagnostics flagged "
        + ", ".join(outliers["CountyName"].tolist())
        + ". These counties deserve sensitivity review; a flag does not mean the data are wrong.",
    )
    set_paragraph(
        doc,
        ("Estimated price is", "The county price target is"),
        "The county price target is based on official recurring advertised offers in 32 counties; the remaining 35 counties use a transparent statewide median estimate because no local public rate was available. No predictor-based or random synthetic price formula is used.",
    )
    set_paragraph(
        doc,
        "Business use requires",
        "Business use requires broader comparable local price coverage, a refreshed Places pull, and direct validation of candidate trade areas.",
    )
    set_paragraph(
        doc,
        "LA Fitness.",
        "LA Fitness. Public online membership signup flow with club-level plan options. Source",
    )
    set_paragraph(
        doc,
        "Planet Fitness.",
        "Planet Fitness. Official membership page with published starting monthly rates. Source",
    )
    set_paragraph(
        doc,
        (
            "Health & Fitness Association.",
            "YouFit, Gold's Gym, and Orangetheory Fitness.",
        ),
        "YouFit, Gold's Gym, and Orangetheory Fitness. Official Florida directories and public enrollment/studio pages. Source",
    )

    set_table(
        doc.tables[0],
        [[
            "Critical limitation: MembershipPrice uses official advertised recurring offers where locally available (32 counties) and a statewide median estimate for 35 counties without a local public rate. It is not a transaction-price survey; plan terms, promotions, and franchise differences limit comparability."
        ]],
    )
    brand_order = [
        "Planet Fitness",
        "Orangetheory Fitness",
        "LA Fitness",
        "Gold's Gym",
        "YouFit",
    ]
    set_table(
        doc.tables[1],
        [["Tracked brand", "Accepted locations"]]
        + [[brand, location_counts.get(brand, 0)] for brand in brand_order],
    )
    stat_rows = [["Variable", "Mean", "Median", "Std. dev.", "Minimum", "Maximum"]]
    currency_variables = {"MembershipPrice", "MedianHouseholdIncome"}
    for variable in summary.index:
        values = []
        for column in ["Mean", "Median", "StdDev", "Min", "Max"]:
            value = float(summary.loc[variable, column])
            values.append(money(value, 2) if variable == "MembershipPrice" else money(value, 0) if variable in currency_variables else f"{value:,.2f}")
        stat_rows.append([variable, *values])
    set_table(doc.tables[2], stat_rows)
    doc.save(path)
    replace_media(
        path,
        [
            FIGURE_DIR / "01_population_vs_price.png",
            FIGURE_DIR / "02_income_vs_price.png",
            FIGURE_DIR / "03_la_locations_vs_price.png",
            FIGURE_DIR / "04_price_histogram.png",
            FIGURE_DIR / "05_correlation_heatmap.png",
        ],
    )


def coefficient_table(frame: pd.DataFrame) -> list[list[object]]:
    rows = [["Variable", "Coefficient", "p-value", "95% confidence interval"]]
    for _, row in frame.iterrows():
        name = "Intercept" if row["Variable"] == "const" else row["Variable"]
        rows.append(
            [
                name,
                f"{row['Coefficient']:.6f}",
                p_value(float(row["pValue"])),
                f"[{row['CI95Lower']:.6f}, {row['CI95Upper']:.6f}]",
            ]
        )
    return rows


def update_predictions_report() -> None:
    path = REPORT_DIR / "Florida_Fitness_Market_Predictions_and_Insights.docx"
    doc = Document(path)
    report = json.loads((OUTPUT_DIR / "report_data.json").read_text(encoding="utf-8"))
    comparison = pd.read_csv(OUTPUT_DIR / "model_comparison.csv")
    full_coefficients = pd.read_csv(OUTPUT_DIR / "regression_full_coefficients.csv")
    refined_coefficients = pd.read_csv(OUTPUT_DIR / "regression_refined_coefficients.csv")
    vif = pd.read_csv(OUTPUT_DIR / "vif.csv")
    diagnostics = pd.read_csv(OUTPUT_DIR / "diagnostic_tests.csv")
    outliers = pd.read_csv(OUTPUT_DIR / "outliers.csv")
    scenarios = pd.read_csv(OUTPUT_DIR / "scenario_predictions.csv")
    rankings = pd.read_csv(OUTPUT_DIR / "expansion_rankings.csv").head(5)

    set_paragraph(
        doc,
        "All 67 Florida counties",
        "All 67 Florida counties | Census + Google Places + official gym pricing | OLS regression",
    )
    set_paragraph(
        doc,
        "Prepared for the MSIS",
        "Prepared for the MSIS Data Science Assignment | September 1, 2026",
    )
    set_paragraph(
        doc,
        (
            "Observed demographic and location data",
            "Observed demographics, locations, and official advertised price observations",
        ),
        "Observed demographics, locations, and official advertised price observations; transparent estimates only where no local offer was available",
    )
    set_paragraph(
        doc,
        ("The analysis combines", "The analysis combines observed Census demographics"),
        "The analysis combines observed Census demographics, 428 Google Places records across 67 Florida counties, and official advertised recurring price observations. Thirty-two counties have local public price observations; the remaining 35 use an explicitly labeled statewide-median estimate. The full six-variable OLS model explains 35.8% of variation in the price target; the refined three-variable model has adjusted R-squared of 0.261.",
    )
    set_paragraph(
        doc,
        "Dependent variable:",
        "Dependent variable: advertised/estimated monthly MembershipPrice target in U.S. dollars.",
    )
    set_paragraph(
        doc,
        "Observed location data:",
        "Observed location data: 428 Google Places records with gym names, formatted street addresses, coordinates, ratings, and review counts; exact-county filtered and deduplicated by Place ID.",
    )
    set_paragraph(
        doc,
        (
            "Where direct county pricing",
            "Price methodology: the long-form pricing file preserves",
        ),
        "Price methodology: the long-form pricing file preserves official plan-level dues, billing frequency, initiation fees, annual fees, promotions, source URLs, and retrieval timestamps for LA Fitness, Planet Fitness, YouFit, Gold's Gym, and Orangetheory Fitness.",
    )
    set_paragraph(
        doc,
        (
            "Price =",
            "For the county target, one lowest recurring advertised offer",
        ),
        "For the county target, one lowest recurring advertised offer was selected per observed local club and averaged within county. Thirty-two counties have local observations; 35 counties use the statewide median of 232 retained local club offers ($43.99). Planet Fitness published starting rates are real but chainwide, so they are not assigned to counties. No predictor-based or random synthetic formula is used.",
    )
    set_paragraph(
        doc,
        "The highest VIF",
        "The highest VIF is 20.42 for CountyPopulation. Values above 5 are flagged because correlated market-size measures can make individual coefficients unstable even when overall fit is useful.",
    )
    set_paragraph(
        doc,
        ("Backward elimination", "Backward elimination removed CompetitorLocations"),
        "Backward elimination removed CompetitorLocations (p = 0.517), AvgGymRating (p = 0.108), and PopulationDensity (p = 0.097). The refined model retains CountyPopulation, LAFitnessLocations, and MedianHouseholdIncome. This mechanical rule is used for the assignment demonstration; real strategy work should also preserve variables supported by business theory.",
    )
    set_paragraph(
        doc,
        (
            "Strongest retained factor:",
            "Strongest retained factor: LAFitnessLocations has the largest absolute",
        ),
        "Strongest retained factor: LAFitnessLocations has the largest absolute one-standard-deviation association (-$13.57; p = 0.0004). This is an association with the price target, not evidence that opening or closing a club causes prices to change.",
    )
    set_paragraph(
        doc,
        ("Competition:", "Competition: the full-model coefficient is"),
        "Competition: the full-model coefficient is -$0.350 per additional tracked competitor (p = 0.517). It was removed during refinement, so the analysis does not establish an independent competition effect after controlling for the other predictors.",
    )
    set_paragraph(
        doc,
        (
            "All interpretations are",
            "All interpretations are associations within a target",
        ),
        "All interpretations are associations within a target that combines local official advertised offers and transparent statewide-median estimates, not causal effects on prices or expansion success.",
    )
    set_paragraph(doc, "Figure 2.", "Figure 2. Observed/estimated county price targets versus refined-model predictions.")
    set_paragraph(
        doc,
        (
            "Each scenario supplies",
            "Each scenario supplies every variable in the original full model",
        ),
        "Each scenario supplies every variable in the original full model. Predictions use the refined model and are extrapolations from a 67-county dataset whose price target is partly estimated. The mean confidence interval describes uncertainty in the expected county price; the wider prediction interval describes uncertainty for a single comparable county outcome.",
    )

    scenario_text = {
        "A - Rural, low competition": "Population 50,000; LA Fitness locations 2; competitors 1; median household income $62,000; density 120.0 per square mile; average rating 4.20.",
        "B - Miami-Dade, high competition": "Population 2,600,000; LA Fitness locations 15; competitors 30; median household income $57,815; density 1,422.1 per square mile; average rating 4.28.",
        "C - Target market: St. Johns County": "Population 273,425; LA Fitness locations 0; competitors 4; median household income $88,794; density 455.2 per square mile; average rating 4.65.",
    }
    set_paragraph(doc, "Population 50,000;", scenario_text["A - Rural, low competition"])
    set_paragraph(doc, "Population 2,600,000;", scenario_text["B - Miami-Dade, high competition"])
    set_paragraph(
        doc,
        ("Population 1,944,375;", "Population 273,425;"),
        scenario_text["C - Target market: St. Johns County"],
    )
    set_paragraph(doc, "C - Target market:", "C - Target market: St. Johns County")
    set_paragraph(
        doc,
        "Figure 5.",
        "Figure 5. Top ten Florida counties by the current market-potential score.",
    )
    recommendation = "Recommendation: Prioritize validation in " + ", ".join(rankings["CountyName"].tolist()[:5]) + ". These are screening priorities, not final real-estate decisions."
    set_table(doc.tables[0], [[recommendation]])
    set_table(
        doc.tables[1],
        [[
            "Decision boundary: The price target uses official local advertised offers in 32 counties and transparent statewide-median estimates in 35 counties. Advertised prices are not transaction prices, and plan/franchise differences limit comparability. Do not commit capital until actual club pricing, leases, trade-area demographics, traffic, and competitor capacity are validated."
        ]],
    )

    model_rows = [["Metric", "Full model", "Refined model"]]
    for label, full_col, refined_col, formatter in [
        ("R-squared", "R2", "R2", lambda value: f"{value:.3f}"),
        ("Adjusted R-squared", "AdjustedR2", "AdjustedR2", lambda value: f"{value:.3f}"),
        ("AIC", "AIC", "AIC", lambda value: f"{value:.2f}"),
        ("BIC", "BIC", "BIC", lambda value: f"{value:.2f}"),
        ("RMSE (USD)", "RMSE", "RMSE", lambda value: money(value, 2)),
    ]:
        full = float(comparison.loc[comparison["Model"].eq("Full"), full_col].iloc[0])
        refined = float(comparison.loc[comparison["Model"].eq("Refined"), refined_col].iloc[0])
        model_rows.append([label, formatter(full), formatter(refined)])
    set_table(doc.tables[2], model_rows)
    set_table(doc.tables[3], coefficient_table(full_coefficients))
    set_table(
        doc.tables[4],
        [["Variable", "VIF", "VIF > 5 assessment"]]
        + [
            [row["Variable"], f"{row['VIF']:.2f}", "Concern" if row["FlagAbove5"] else "Acceptable"]
            for _, row in vif.iterrows()
        ],
    )
    set_table(doc.tables[5], coefficient_table(refined_coefficients))
    set_table(
        doc.tables[6],
        [["Diagnostic test", "Statistic", "p-value", "Interpretation"]]
        + [
            [row["Test"], f"{row['Statistic']:.3f}", p_value(float(row["pValue"])), row["Interpretation"]]
            for _, row in diagnostics.iterrows()
        ],
    )
    set_table(
        doc.tables[7],
        [["County", "Studentized residual", "Cook's distance", "Leverage"]]
        + [
            [row["CountyName"], f"{row['StudentizedResidual']:.2f}", f"{row['CooksDistance']:.3f}", f"{row['Leverage']:.3f}"]
            for _, row in outliers.iterrows()
        ],
    )
    scenario_rows = [["Scenario", "Predicted price", "95% mean CI", "95% prediction interval"]]
    for _, row in scenarios.iterrows():
        scenario_rows.append(
            [
                row["Scenario"],
                money(row["PredictedPrice"]),
                f"{money(row['MeanCI95Lower'])} - {money(row['MeanCI95Upper'])}",
                f"{money(row['PredictionInterval95Lower'])} - {money(row['PredictionInterval95Upper'])}",
            ]
        )
    set_table(doc.tables[8], scenario_rows)
    ranking_rows = [["Rank", "County", "Score", "Population", "Median income", "Predicted price"]]
    for _, row in rankings.iterrows():
        ranking_rows.append(
            [
                int(row["Rank"]),
                row["CountyName"],
                f"{row['MarketPotentialScore']:.1f}",
                f"{int(row['CountyPopulation']):,}",
                money(row["MedianHouseholdIncome"], 0),
                money(row["PredictedMembershipPrice"]),
            ]
        )
    set_table(doc.tables[9], ranking_rows)

    for rank, (_, row) in enumerate(rankings.iterrows(), start=1):
        old_names = ["Broward County", "St. Johns County", "Miami-Dade County", "Orange County", "Palm Beach County"]
        prefix = (f"{rank}. {old_names[rank - 1]}", f"{rank}. {row['CountyName']}")
        strengths = "large demand base" if row["CountyPopulation"] >= 500000 else "favorable purchasing power"
        if row["LAFitnessLocations"] == 0:
            strengths += ", LA Fitness whitespace"
        else:
            strengths += ", existing market evidence"
        if row["CompetitorLocations"] <= 5:
            strengths += ", light tracked-competitor intensity"
        paragraph_text = (
            f"{rank}. {row['CountyName']} | Score {row['MarketPotentialScore']:.1f}/100. "
            f"Key screening strengths: {strengths}. Observed LA Fitness locations: {int(row['LAFitnessLocations'])}; "
            f"tracked competitors: {int(row['CompetitorLocations'])}."
        )
        set_paragraph(doc, prefix, paragraph_text)

    score_paragraphs = [
        paragraph
        for paragraph in doc.paragraphs
        if paragraph.text.strip().startswith("Score ")
    ]
    if len(score_paragraphs) != len(rankings):
        raise ValueError("Could not find the five ranking-detail paragraphs in the report")
    for paragraph, (_, row) in zip(score_paragraphs, rankings.iterrows()):
        strengths = "large demand base" if row["CountyPopulation"] >= 500000 else "favorable purchasing power"
        if row["LAFitnessLocations"] == 0:
            strengths += ", LA Fitness whitespace"
        else:
            strengths += ", existing market evidence"
        if row["CompetitorLocations"] <= 5:
            strengths += ", light tracked-competitor intensity"
        paragraph.text = (
            f"Score {row['MarketPotentialScore']:.1f}/100. Key screening strengths: {strengths}. "
            f"Observed LA Fitness locations: {int(row['LAFitnessLocations'])}; "
            f"tracked competitors: {int(row['CompetitorLocations'])}."
        )

    set_paragraph(doc, "Validate current club rosters", "Validate current club rosters, closures, planned openings, and public price offers in the top five counties.")
    set_paragraph(doc, "Collect observed monthly dues", "Collect observed monthly dues, initiation fees, annual fees, promotions, and contract terms for a representative sample of clubs.")
    set_paragraph(doc, "Replace county boundaries", "Replace county boundaries with 10- to 15-minute drive-time trade areas around candidate sites.")
    set_paragraph(doc, "Add lease rates", "Add lease rates, daytime population, age mix, traffic counts, and competitor square footage before site underwriting.")
    set_paragraph(doc, "Re-estimate the model", "Re-estimate the model with broader observed comparable prices and validate out of sample before using it for capital allocation.")
    set_paragraph(doc, ("Estimated target:", "Price target coverage:"), "Price target coverage: 32 counties use local official advertised offers; 35 use the transparent statewide median because no local public rate was available. The target is not a transaction-price survey.")
    set_paragraph(doc, "Search coverage:", "Search coverage: Google Places Text Search ranks candidates and does not guarantee a complete chain inventory.")
    set_paragraph(doc, "Temporal mismatch:", "Temporal mismatch: demographics use 2020 population and 2021 income, while locations and prices are retrieval-time snapshots.")
    set_paragraph(doc, "Aggregation:", "Aggregation: county averages mask neighborhood trade areas and cross-county customer travel.")
    set_paragraph(doc, "Ranking sensitivity:", "Ranking sensitivity: market-potential results depend on subjective weights, min-max scaling, and the observed-versus-estimated price mix.")
    set_paragraph(
        doc,
        ("fitness_market_analysis.py:", "scripts/fitness_market_analysis.py:"),
        "scripts/fitness_market_analysis.py: API collection, official-price integration, EDA, regression, diagnostics, predictions, and rankings.",
    )
    set_paragraph(doc, "data/MainDataset.csv:", "data/MainDataset.csv: 67 rows and 8 required columns, rebuilt from the current observed/estimated price target.")
    set_paragraph(doc, "outputs/:", "outputs/: model tables, diagnostic tables, pricing-coverage audit, scenario predictions, rankings, and current figures.")
    doc.save(path)
    replace_media(
        path,
        [
            FIGURE_DIR / "09_standardized_effects.png",
            FIGURE_DIR / "06_actual_vs_predicted.png",
            FIGURE_DIR / "07_qq_plot.png",
            FIGURE_DIR / "08_residuals_vs_fitted.png",
            FIGURE_DIR / "10_top_markets.png",
        ],
    )


if __name__ == "__main__":
    update_eda_report()
    update_predictions_report()
    print("Updated both Word reports from the current analysis outputs.")
