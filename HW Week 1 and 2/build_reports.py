#!/usr/bin/env python3
"""Build the two Word reports required by the fitness-market assignment."""

from __future__ import annotations

import json
import sys
from datetime import date
from pathlib import Path
from typing import Iterable, Sequence

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
DATA_DIR = ROOT / "data"
OUTPUT_DIR = ROOT / "outputs"
FIGURE_DIR = OUTPUT_DIR / "figures"
EDA_PATH = ROOT / "Florida_Fitness_Market_EDA_Report.docx"
INSIGHTS_PATH = ROOT / "Florida_Fitness_Market_Predictions_and_Insights.docx"

DOC_SKILL = Path(
    "/Users/chris/.codex-personal/plugins/cache/openai-primary-runtime/"
    "documents/26.826.12353/skills/documents"
)
sys.path.insert(0, str(DOC_SKILL / "scripts"))
from table_geometry import apply_table_geometry, exact_column_widths  # noqa: E402


NAVY = RGBColor(18, 53, 91)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
TEAL = RGBColor(42, 157, 143)
GOLD = RGBColor(217, 154, 43)
INK = RGBColor(38, 55, 70)
GRAY = RGBColor(92, 104, 116)
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "EEF3F7"
LIGHT_GOLD = "FFF7E6"
WHITE = RGBColor(255, 255, 255)
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120

REFERENCE_SOURCES = [
    (
        "U.S. Census Bureau. 2020 Census Redistricting Data, Table P1, variable P1_001N.",
        "https://api.census.gov/data/2020/dec/pl/groups/P1.html",
    ),
    (
        "U.S. Census Bureau. 2021 ACS 5-Year Estimates, Table B19013, median household income.",
        "https://data.census.gov/table/ACSDT5Y2021.B19013",
    ),
    (
        "U.S. Census Bureau. 2020 Gazetteer Files and record layout.",
        "https://www.census.gov/geographies/reference-files/2020/geo/gazetter-file.html",
    ),
    (
        "Google Maps Platform. Text Search (New) documentation and pagination requirements.",
        "https://developers.google.com/maps/documentation/places/web-service/text-search",
    ),
    (
        "LA Fitness. Public online membership rate example with monthly plan options.",
        "https://www.lafitness.com/Pages/MembershipSignUpRate.aspx?id=b0J8ga4R865EceAG6%2FCX9A%3D%3D",
    ),
    (
        "Planet Fitness. Official membership page with starting monthly rates.",
        "https://www.planetfitness.com/gym-memberships",
    ),
    (
        "Health & Fitness Association. U.S. facility membership and average-dues overview.",
        "https://www.healthandfitness.org/u-s-fitness-facility-memberships-reach-the-highest-level-ever-as-dues-rise/",
    ),
]


def set_run_font(
    run,
    *,
    name: str = "Calibri",
    size: float | None = None,
    color: RGBColor | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instruction, separate, text, end):
        run._r.append(element)
    set_run_font(run, size=8.5, color=GRAY)


def add_hyperlink(paragraph, label: str, url: str) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2E74B5")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend([color, underline])
    text = OxmlElement("w:t")
    text.text = label
    run.extend([run_properties, text])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def configure_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    caption = styles["Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    caption.font.size = Pt(9)
    caption.font.color.rgb = GRAY
    caption.font.italic = True
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(9)
    caption.paragraph_format.keep_together = True

    set_running_furniture(doc)
    return doc


def set_running_furniture(doc: Document) -> None:
    section = doc.sections[0]
    doc.settings.odd_and_even_pages_header_footer = True

    for header in (section.header, section.even_page_header):
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("MSIS DATA SCIENCE  |  FLORIDA FITNESS MARKET")
        set_run_font(run, size=8, color=GRAY, bold=True)

    for footer in (section.footer, section.even_page_footer):
        p = footer.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.25))
        left = p.add_run("Academic analysis - estimated county price target")
        set_run_font(left, size=8, color=GRAY)
        p.add_run("\t")
        add_page_field(p)


def add_cover(doc: Document, report_label: str, title: str, subtitle: str) -> None:
    for _ in range(3):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(16)
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run(report_label.upper())
    set_run_font(run, size=10.5, color=GOLD, bold=True)

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(9)
    run = title_p.add_run(title)
    set_run_font(run, size=29, color=NAVY, bold=True)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(24)
    run = subtitle_p.add_run(subtitle)
    set_run_font(run, size=14, color=DARK_BLUE)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(6)
    run = meta.add_run("All 67 Florida counties | Census + Google Places | OLS regression")
    set_run_font(run, size=10.5, color=GRAY, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("Prepared for the MSIS Data Science Assignment | August 29, 2026")
    set_run_font(run, size=10, color=GRAY, italic=True)

    for _ in range(3):
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(16)
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(12)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run(
        "Observed demographic and location data; transparently estimated membership-price target"
    )
    set_run_font(run, size=9.5, color=GRAY, italic=True)
    doc.add_page_break()


def add_callout(doc: Document, label: str, text: str, fill: str = LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_repeat_table_header(table.rows[0])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    label_run = p.add_run(f"{label}: ")
    set_run_font(label_run, size=10.5, color=NAVY, bold=True)
    text_run = p.add_run(text)
    set_run_font(text_run, size=10.5, color=INK)
    apply_table_geometry(
        table,
        [TABLE_WIDTH_DXA],
        table_width_dxa=TABLE_WIDTH_DXA,
        indent_dxa=TABLE_INDENT_DXA,
    )
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(3)


def add_bullet(doc: Document, text: str, *, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        set_run_font(run, bold=True, color=NAVY)
        p.add_run(text[len(bold_prefix) :])
    else:
        p.add_run(text)


def add_numbered(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def format_value(value: object, kind: str = "text") -> str:
    if value is None or pd.isna(value):
        return "-"
    number = float(value) if kind != "text" else None
    if kind == "integer":
        return f"{number:,.0f}"
    if kind == "decimal1":
        return f"{number:,.1f}"
    if kind == "decimal2":
        return f"{number:,.2f}"
    if kind == "currency":
        return f"${number:,.2f}"
    if kind == "currency0":
        return f"${number:,.0f}"
    if kind == "percent":
        return f"{number:.1%}"
    if kind == "pvalue":
        return "<0.001" if number < 0.001 else f"{number:.3f}"
    return str(value)


def add_data_table(
    doc: Document,
    headers: Sequence[str],
    rows: Iterable[Sequence[object]],
    widths_dxa: Sequence[int],
    *,
    alignments: Sequence[WD_ALIGN_PARAGRAPH] | None = None,
    font_size: float = 9.2,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.rows[0].cells
    for column, header in enumerate(headers):
        cell = table.rows[0].cells[column]
        set_cell_shading(cell, LIGHT_GRAY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = (
            alignments[column] if alignments else WD_ALIGN_PARAGRAPH.CENTER
        )
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(header)
        set_run_font(run, size=font_size, color=NAVY, bold=True)
    set_repeat_table_header(table.rows[0])

    for row_values in rows:
        cells = table.add_row().cells
        for column, value in enumerate(row_values):
            cell = cells[column]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = (
                alignments[column]
                if alignments
                else (WD_ALIGN_PARAGRAPH.LEFT if column == 0 else WD_ALIGN_PARAGRAPH.RIGHT)
            )
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(str(value))
            set_run_font(run, size=font_size, color=INK)

    widths = exact_column_widths(widths_dxa, TABLE_WIDTH_DXA)
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=TABLE_WIDTH_DXA,
        indent_dxa=TABLE_INDENT_DXA,
    )
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(4)
    after.paragraph_format.space_after = Pt(0)


def add_figure(
    doc: Document,
    filename: str,
    caption: str,
    alt_text: str,
    *,
    width: float = 6.15,
    page_break_before: bool = False,
) -> None:
    image_p = doc.add_paragraph()
    image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_p.paragraph_format.page_break_before = page_break_before
    image_p.paragraph_format.space_before = Pt(4)
    image_p.paragraph_format.space_after = Pt(0)
    image_p.paragraph_format.keep_with_next = True
    run = image_p.add_run()
    shape = run.add_picture(str(FIGURE_DIR / filename), width=Inches(width))
    shape._inline.docPr.set("descr", alt_text)
    caption_p = doc.add_paragraph(caption, style="Caption")
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_sources(doc: Document, *, page_break_before: bool = False) -> None:
    heading = doc.add_heading("References", level=1)
    heading.paragraph_format.page_break_before = page_break_before
    for label, url in REFERENCE_SOURCES:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.first_line_indent = Inches(-0.22)
        p.paragraph_format.space_after = Pt(5)
        text_run = p.add_run(label + " ")
        set_run_font(text_run, size=9.5, color=INK)
        add_hyperlink(p, "Source", url)


def load_inputs() -> dict[str, object]:
    summary = pd.read_csv(OUTPUT_DIR / "summary_statistics.csv")
    summary = summary.rename(columns={summary.columns[0]: "Variable"})
    data: dict[str, object] = {
        "main": pd.read_csv(DATA_DIR / "MainDataset.csv"),
        "locations": pd.read_csv(DATA_DIR / "gym_locations.csv"),
        "sources": pd.read_csv(DATA_DIR / "data_sources.csv"),
        "summary": summary,
        "correlation": pd.read_csv(OUTPUT_DIR / "correlation_matrix.csv", index_col=0),
        "full_coefficients": pd.read_csv(
            OUTPUT_DIR / "regression_full_coefficients.csv"
        ),
        "refined_coefficients": pd.read_csv(
            OUTPUT_DIR / "regression_refined_coefficients.csv"
        ),
        "vif": pd.read_csv(OUTPUT_DIR / "vif.csv"),
        "comparison": pd.read_csv(OUTPUT_DIR / "model_comparison.csv"),
        "diagnostics": pd.read_csv(OUTPUT_DIR / "diagnostic_tests.csv"),
        "outliers": pd.read_csv(OUTPUT_DIR / "outliers.csv"),
        "rankings": pd.read_csv(OUTPUT_DIR / "expansion_rankings.csv"),
        "scenarios": pd.read_csv(OUTPUT_DIR / "scenario_predictions.csv"),
        "effects": pd.read_csv(OUTPUT_DIR / "standardized_effects.csv"),
        "quality": pd.read_csv(OUTPUT_DIR / "data_quality_checks.csv"),
    }
    data["payload"] = json.loads(
        (OUTPUT_DIR / "report_data.json").read_text(encoding="utf-8")
    )
    return data


def build_eda_report(data: dict[str, object]) -> None:
    main: pd.DataFrame = data["main"]
    locations: pd.DataFrame = data["locations"]
    summary: pd.DataFrame = data["summary"]
    correlation: pd.DataFrame = data["correlation"]
    outliers: pd.DataFrame = data["outliers"]
    quality: pd.DataFrame = data["quality"]

    doc = configure_document()
    add_cover(
        doc,
        "Exploratory Data Analysis Report",
        "Fitness Market Analysis",
        "LA Fitness Expansion Strategy Across Florida Counties",
    )
    doc.add_heading("Purpose and Scope", level=1)
    doc.add_paragraph(
        "This report examines the market structure of all 67 Florida counties before "
        "regression modeling. It combines official demographic measures with live Google "
        "Places results for LA Fitness and four competitor brands. The objective is to "
        "identify distributions, relationships, and unusual counties that may affect model "
        "interpretation or expansion decisions."
    )
    add_callout(
        doc,
        "Critical limitation",
        "MembershipPrice is an estimated teaching variable, not an observed county price "
        "survey. It is generated from documented market inputs and random noise (seed 6643). "
        "EDA findings involving price demonstrate the workflow and must not be treated as "
        "independent market evidence.",
        LIGHT_GOLD,
    )

    doc.add_heading("Data Collection and Preparation", level=1)
    doc.add_heading("Observed demographic inputs", level=2)
    add_bullet(doc, "CountyPopulation: 2020 Census PL 94-171 total population (P1_001N).")
    add_bullet(doc, "MedianHouseholdIncome: 2021 ACS five-year estimate (B19013_001E).")
    add_bullet(doc, "LandAreaSqMiles: 2020 Census Gazetteer; PopulationDensity is population divided by land area.")

    doc.add_heading("Observed gym-market inputs", level=2)
    chain_counts = (
        locations.groupby("GymChain").size().sort_values(ascending=False).reset_index(name="Locations")
    )
    doc.add_paragraph(
        f"The Google Places Text Search pipeline issued branded, county-specific queries, "
        f"filtered results to exact county and brand matches, and deduplicated by Place ID. "
        f"It retained {len(locations):,} locations across the five tracked brands."
    )
    add_data_table(
        doc,
        ["Tracked brand", "Accepted locations"],
        [(row.GymChain, f"{row.Locations:,}") for row in chain_counts.itertuples()],
        [6500, 2860],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT],
        font_size=9.6,
    )
    doc.add_paragraph(
        "Google Places results are a retrieval-time snapshot. Search relevance, closures, "
        "brand naming, and data coverage can change, so counts should be validated before a "
        "real estate commitment."
    )

    heading = doc.add_heading("Missing-Data Treatment", level=2)
    heading.paragraph_format.page_break_before = True
    imputed_rating_count = int(
        float(
            quality.loc[
                quality["Check"].eq("Counties with imputed average rating"), "Result"
            ].iloc[0]
        )
    )
    add_bullet(doc, "Missing LA Fitness and competitor counts were set to zero after the county merge.")
    add_bullet(
        doc,
        f"{imputed_rating_count} counties with no retained tracked-brand ratings received the statewide median Google "
        "Places rating. The completed MainDataset contains no missing values.",
    )
    add_bullet(doc, "Post-integration missing values in the analytical dataset: 0.")

    doc.add_heading("Descriptive Statistics", level=1)
    statistic_rows = []
    for row in summary.itertuples(index=False):
        variable = row.Variable
        currency_variables = {"MembershipPrice", "MedianHouseholdIncome"}
        formatter = "currency0" if variable == "MedianHouseholdIncome" else "currency" if variable == "MembershipPrice" else "decimal2"
        statistic_rows.append(
            [
                variable,
                format_value(row.Mean, formatter),
                format_value(row.Median, formatter),
                format_value(row.StdDev, formatter),
                format_value(row.Min, formatter),
                format_value(row.Max, formatter),
            ]
        )
    add_data_table(
        doc,
        ["Variable", "Mean", "Median", "Std. dev.", "Minimum", "Maximum"],
        statistic_rows,
        [2550, 1362, 1362, 1362, 1362, 1362],
        font_size=8.8,
    )

    doc.add_heading("Required Visual Analysis", level=1)
    add_figure(
        doc,
        "01_population_vs_price.png",
        "Figure 1. County population versus estimated monthly membership price.",
        "Scatter plot of Florida county population against estimated membership price with a fitted trend line.",
    )
    add_figure(
        doc,
        "02_income_vs_price.png",
        "Figure 2. Median household income versus estimated monthly membership price.",
        "Scatter plot of county median household income against estimated membership price.",
        width=5.2,
        page_break_before=True,
    )
    add_figure(
        doc,
        "03_la_locations_vs_price.png",
        "Figure 3. Observed LA Fitness location count versus estimated price.",
        "Scatter plot of observed LA Fitness locations by county against estimated membership price.",
        width=5.2,
    )
    add_figure(
        doc,
        "04_price_histogram.png",
        "Figure 4. Distribution of estimated county membership prices.",
        "Histogram of estimated county membership price with the mean marked.",
    )
    add_figure(
        doc,
        "05_correlation_heatmap.png",
        "Figure 5. Pearson correlation matrix for all analytical variables.",
        "Lower-triangle heatmap of correlations among price and six independent variables.",
        width=5.5,
        page_break_before=True,
    )

    doc.add_heading("Correlation Findings", level=1)
    price_corr = correlation["MembershipPrice"].drop("MembershipPrice").sort_values(
        key=lambda s: s.abs(), ascending=False
    )
    for variable, value in price_corr.items():
        direction = "positive" if value >= 0 else "negative"
        add_bullet(doc, f"{variable}: r = {value:.3f} ({direction}).")
    doc.add_paragraph(
        "The market-size variables are also related to one another. This is expected because "
        "large, dense, affluent counties tend to support more fitness locations. Those "
        "relationships motivate the formal VIF check in the regression report."
    )

    heading = doc.add_heading("Outliers and Business Meaning", level=1)
    heading.paragraph_format.page_break_before = True
    population_max = main.loc[main["CountyPopulation"].idxmax()]
    density_max = main.loc[main["PopulationDensity"].idxmax()]
    income_max = main.loc[main["MedianHouseholdIncome"].idxmax()]
    add_bullet(
        doc,
        f"Population: {population_max['CountyName']} is the largest county market at "
        f"{population_max['CountyPopulation']:,.0f} residents, giving it disproportionate "
        "leverage in raw population relationships.",
    )
    add_bullet(
        doc,
        f"Density: {density_max['CountyName']} has the highest density at "
        f"{density_max['PopulationDensity']:,.1f} residents per square mile, reflecting a "
        "different site-selection environment from geographically large rural counties.",
    )
    add_bullet(
        doc,
        f"Purchasing power: {income_max['CountyName']} has the highest 2021 median household "
        f"income at ${income_max['MedianHouseholdIncome']:,.0f}.",
    )
    if not outliers.empty:
        add_bullet(
            doc,
            "Influence diagnostics flagged "
            + ", ".join(outliers["CountyName"].head(6).tolist())
            + ". These counties deserve sensitivity review; a flag does not mean the data are wrong.",
        )

    doc.add_heading("EDA Conclusions", level=1)
    add_numbered(doc, "County markets differ sharply in size, density, purchasing power, and chain presence.")
    add_numbered(doc, "Estimated price is most strongly associated with a subset of the market variables, but the target is synthetic by design.")
    add_numbered(doc, "The 67-row dataset is complete and suitable for the required OLS demonstration, subject to multicollinearity and influence checks.")
    add_numbered(doc, "Business use requires observed county pricing, a refreshed Places pull, and direct validation of candidate trade areas.")
    add_sources(doc, page_break_before=True)
    doc.save(EDA_PATH)


def build_insights_report(data: dict[str, object]) -> None:
    main: pd.DataFrame = data["main"]
    full_coefficients: pd.DataFrame = data["full_coefficients"]
    refined_coefficients: pd.DataFrame = data["refined_coefficients"]
    vif: pd.DataFrame = data["vif"]
    comparison: pd.DataFrame = data["comparison"]
    diagnostics: pd.DataFrame = data["diagnostics"]
    outliers: pd.DataFrame = data["outliers"]
    rankings: pd.DataFrame = data["rankings"]
    scenarios: pd.DataFrame = data["scenarios"]
    effects: pd.DataFrame = data["effects"]
    payload: dict[str, object] = data["payload"]

    doc = configure_document()
    add_cover(
        doc,
        "Predictions & Business Insights Report",
        "LA Fitness Expansion Strategy",
        "Regression Results, Scenario Pricing, and Florida Market Priorities",
    )

    doc.add_heading("Executive Summary", level=1)
    full = payload["models"]["full"]
    refined = payload["models"]["refined"]
    top_five = rankings.head(5)
    doc.add_paragraph(
        f"The analysis combines observed Census demographics with {payload['dataset']['location_rows']:,} "
        "Google Places records across 67 Florida counties. The full six-variable OLS model "
        f"explains {full['r2']:.1%} of variation in the estimated price target; the refined "
        f"model retains {', '.join(refined['variables'])} and achieves adjusted R-squared "
        f"of {refined['adjusted_r2']:.3f}."
    )
    add_callout(
        doc,
        "Recommendation",
        f"Prioritize validation in {top_five.iloc[0]['CountyName']}, then "
        f"{top_five.iloc[1]['CountyName']}, {top_five.iloc[2]['CountyName']}, "
        f"{top_five.iloc[3]['CountyName']}, and {top_five.iloc[4]['CountyName']}. "
        "These are screening priorities, not final real-estate decisions.",
    )
    add_callout(
        doc,
        "Decision boundary",
        "The membership-price target is estimated, so coefficient significance and county "
        "rankings demonstrate an analytical workflow rather than independently verified "
        "economics. Do not commit capital until actual club pricing, leases, trade-area "
        "demographics, traffic, and competitor capacity are validated.",
        LIGHT_GOLD,
    )

    heading = doc.add_heading("Data and Modeling Approach", level=1)
    heading.paragraph_format.page_break_before = True
    add_bullet(doc, "Unit of analysis: one row per Florida county (n = 67).")
    add_bullet(doc, "Dependent variable: estimated average monthly MembershipPrice in U.S. dollars.")
    add_bullet(doc, "Independent variables: population, LA Fitness count, competitor count, household income, density, and average rating.")
    add_bullet(doc, "Observed location data: live Google Places Text Search results, exact-county filtered and deduplicated by Place ID.")
    add_bullet(doc, "Estimation method: Ordinary Least Squares using Python statsmodels.")
    doc.add_paragraph(
        "Where direct county pricing was unavailable, the assignment permits simulated data. "
        "The price target uses a documented teaching formula anchored to public chain and "
        "industry price points, plus normally distributed noise. The formula was clipped to "
        "$24-$55 per month and should not be interpreted as a survey:"
    )
    formula = doc.add_paragraph()
    formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
    formula.paragraph_format.space_before = Pt(5)
    formula.paragraph_format.space_after = Pt(9)
    run = formula.add_run(
        "Price = 8 + 0.0000015(Population) + 0.45(LA) - 0.12(Competitors) "
        "+ 0.00012(Income) + 0.001(Density) + 5(Rating) + error"
    )
    set_run_font(run, name="Courier New", size=8.8, color=NAVY, bold=True)

    heading = doc.add_heading("Full OLS Model", level=1)
    heading.paragraph_format.page_break_before = True
    add_data_table(
        doc,
        ["Metric", "Full model", "Refined model"],
        [
            ("R-squared", f"{comparison.iloc[0]['R2']:.3f}", f"{comparison.iloc[1]['R2']:.3f}"),
            ("Adjusted R-squared", f"{comparison.iloc[0]['AdjustedR2']:.3f}", f"{comparison.iloc[1]['AdjustedR2']:.3f}"),
            ("AIC", f"{comparison.iloc[0]['AIC']:.2f}", f"{comparison.iloc[1]['AIC']:.2f}"),
            ("BIC", f"{comparison.iloc[0]['BIC']:.2f}", f"{comparison.iloc[1]['BIC']:.2f}"),
            ("RMSE (USD)", f"${comparison.iloc[0]['RMSE']:.2f}", f"${comparison.iloc[1]['RMSE']:.2f}"),
        ],
        [3900, 2730, 2730],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT],
        font_size=9.4,
    )

    full_rows = []
    for row in full_coefficients.itertuples(index=False):
        label = "Intercept" if row.Variable == "const" else row.Variable
        full_rows.append(
            [
                label,
                f"{row.Coefficient:,.6f}",
                format_value(row.pValue, "pvalue"),
                f"[{row.CI95Lower:,.4f}, {row.CI95Upper:,.4f}]",
            ]
        )
    add_data_table(
        doc,
        ["Variable", "Coefficient", "p-value", "95% confidence interval"],
        full_rows,
        [2800, 1800, 1300, 3460],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.CENTER],
        font_size=8.8,
    )

    heading = doc.add_heading("Multicollinearity", level=2)
    heading.paragraph_format.page_break_before = True
    vif_rows = [
        [row.Variable, f"{row.VIF:.2f}", "Concern" if bool(row.FlagAbove5) else "Acceptable"]
        for row in vif.itertuples(index=False)
    ]
    add_data_table(
        doc,
        ["Variable", "VIF", "VIF > 5 assessment"],
        vif_rows,
        [4750, 1500, 3110],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.CENTER],
        font_size=9.2,
    )
    max_vif_row = vif.loc[vif["VIF"].idxmax()]
    doc.add_paragraph(
        f"The highest VIF is {max_vif_row['VIF']:.2f} for {max_vif_row['Variable']}. "
        "Values above 5 are flagged because correlated market-size measures can make "
        "individual coefficients unstable even when overall fit is useful."
    )

    heading = doc.add_heading("Refined Model", level=1)
    heading.paragraph_format.page_break_before = True
    removed = [entry["RemovedVariable"] for entry in payload["models"]["elimination_log"]]
    doc.add_paragraph(
        "Backward elimination removed variables with p-values above 0.05, one at a time. "
        + ("Removed variables: " + ", ".join(removed) + "." if removed else "No variables were removed.")
        + " This mechanical rule was used for the assignment demonstration; real strategy "
        "work should also preserve variables supported by business theory."
    )
    refined_rows = []
    for row in refined_coefficients.itertuples(index=False):
        label = "Intercept" if row.Variable == "const" else row.Variable
        refined_rows.append(
            [
                label,
                f"{row.Coefficient:,.6f}",
                format_value(row.pValue, "pvalue"),
                f"[{row.CI95Lower:,.4f}, {row.CI95Upper:,.4f}]",
            ]
        )
    add_data_table(
        doc,
        ["Variable", "Coefficient", "p-value", "95% confidence interval"],
        refined_rows,
        [2800, 1800, 1300, 3460],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.CENTER],
        font_size=8.8,
    )
    add_figure(
        doc,
        "09_standardized_effects.png",
        "Figure 1. Refined-model price effect associated with a one-standard-deviation increase in each retained predictor.",
        "Horizontal bar chart of standardized dollar effects for retained regression variables.",
        width=6.0,
    )

    strongest = effects.loc[effects["DollarEffectPerOneSD"].abs().idxmax()]
    competition_row = full_coefficients.loc[
        full_coefficients["Variable"].eq("CompetitorLocations")
    ].iloc[0]
    doc.add_heading("Interpretation", level=2)
    add_bullet(
        doc,
        f"Strongest retained factor: {strongest['Variable']} has the largest absolute "
        f"one-standard-deviation effect (${abs(strongest['DollarEffectPerOneSD']):.2f}).",
    )
    add_bullet(
        doc,
        f"Competition: the full-model coefficient is {competition_row['Coefficient']:.3f} "
        f"dollars per additional tracked competitor (p = {competition_row['pValue']:.3f}). "
        "It was removed during refinement, so the analysis does not establish an independent "
        "competition effect after controlling for the other predictors.",
    )
    add_bullet(doc, "All interpretations are associations within an estimated teaching target, not causal effects.")

    heading = doc.add_heading("Model Diagnostics", level=1)
    heading.paragraph_format.page_break_before = True
    diagnostic_rows = [
        [row.Test, f"{row.Statistic:.3f}", format_value(row.pValue, "pvalue"), row.Interpretation]
        for row in diagnostics.itertuples(index=False)
    ]
    add_data_table(
        doc,
        ["Diagnostic test", "Statistic", "p-value", "Interpretation"],
        diagnostic_rows,
        [3400, 1200, 1200, 3560],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.LEFT],
        font_size=8.8,
    )
    add_figure(
        doc,
        "06_actual_vs_predicted.png",
        "Figure 2. Actual estimated prices versus refined-model predictions.",
        "Scatter plot comparing estimated price targets to refined model predictions with a 45-degree line.",
        width=5.9,
    )
    add_figure(
        doc,
        "07_qq_plot.png",
        "Figure 3. Q-Q plot of refined-model residuals.",
        "Quantile-quantile plot used to evaluate residual normality.",
        width=5.3,
    )
    add_figure(
        doc,
        "08_residuals_vs_fitted.png",
        "Figure 4. Refined-model residuals versus fitted values.",
        "Scatter plot of residuals against fitted membership prices with a zero reference line.",
        width=5.3,
    )
    if not outliers.empty:
        heading = doc.add_heading("Influential Observations", level=2)
        heading.paragraph_format.page_break_before = True
        outlier_rows = [
            [
                row.CountyName,
                f"{row.StudentizedResidual:.2f}",
                f"{row.CooksDistance:.3f}",
                f"{row.Leverage:.3f}",
            ]
            for row in outliers.head(8).itertuples(index=False)
        ]
        add_data_table(
            doc,
            ["County", "Studentized residual", "Cook's distance", "Leverage"],
            outlier_rows,
            [3600, 1920, 1920, 1920],
            alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT],
            font_size=9.0,
        )

    doc.add_heading("Scenario Predictions", level=1)
    doc.add_paragraph(
        "Each scenario supplies every variable in the original full model. Predictions use "
        "the refined model. The mean confidence interval describes uncertainty in the "
        "expected county price; the wider prediction interval describes uncertainty for a "
        "single comparable county outcome."
    )
    scenario_rows = []
    for row in scenarios.itertuples(index=False):
        scenario_rows.append(
            [
                row.Scenario,
                f"${row.PredictedPrice:.2f}",
                f"${row.MeanCI95Lower:.2f} - ${row.MeanCI95Upper:.2f}",
                f"${row.PredictionInterval95Lower:.2f} - ${row.PredictionInterval95Upper:.2f}",
            ]
        )
    add_data_table(
        doc,
        ["Scenario", "Predicted price", "95% mean CI", "95% prediction interval"],
        scenario_rows,
        [3500, 1600, 2000, 2260],
        alignments=[WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER],
        font_size=8.8,
    )
    for row in scenarios.itertuples(index=False):
        doc.add_heading(row.Scenario, level=2)
        doc.add_paragraph(
            f"Population {row.CountyPopulation:,.0f}; LA Fitness locations {row.LAFitnessLocations:,.0f}; "
            f"competitors {row.CompetitorLocations:,.0f}; median household income "
            f"${row.MedianHouseholdIncome:,.0f}; density {row.PopulationDensity:,.1f} per "
            f"square mile; average rating {row.AvgGymRating:.2f}."
        )

    heading = doc.add_heading("Expansion Recommendations", level=1)
    heading.paragraph_format.page_break_before = True
    doc.add_paragraph(
        "The market-potential score is a transparent screening index calculated within this "
        "67-county dataset. Weights are 30% demand (log population), 25% purchasing power, "
        "20% low competitor intensity, 15% predicted price, and 10% low LA Fitness intensity. "
        "Components are min-max normalized, so the score is relative and not transferable to "
        "another state or vintage without recalculation."
    )
    top_rows = []
    for row in rankings.head(5).itertuples(index=False):
        top_rows.append(
            [
                int(row.Rank),
                row.CountyName,
                f"{row.MarketPotentialScore:.1f}",
                f"{row.CountyPopulation:,.0f}",
                f"${row.MedianHouseholdIncome:,.0f}",
                f"${row.PredictedMembershipPrice:.2f}",
            ]
        )
    add_data_table(
        doc,
        ["Rank", "County", "Score", "Population", "Median income", "Predicted price"],
        top_rows,
        [700, 2350, 1100, 1700, 1750, 1760],
        alignments=[WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.RIGHT],
        font_size=8.6,
    )
    add_figure(
        doc,
        "10_top_markets.png",
        "Figure 5. Top ten Florida counties by the simulated market-potential score.",
        "Horizontal bar chart ranking the top ten counties by market potential score.",
        width=5.55,
    )
    for row in rankings.head(5).itertuples(index=False):
        doc.add_heading(f"{int(row.Rank)}. {row.CountyName}", level=2)
        strengths = []
        if row.DemandComponent >= 0.75:
            strengths.append("large demand base")
        if row.PurchasingPowerComponent >= 0.75:
            strengths.append("strong purchasing power")
        if row.LowCompetitionComponent >= 0.65:
            strengths.append("favorable tracked-competitor intensity")
        if row.WhitespaceComponent >= 0.65:
            strengths.append("LA Fitness whitespace")
        if not strengths:
            strengths.append("balanced performance across the screening components")
        doc.add_paragraph(
            f"Score {row.MarketPotentialScore:.1f}/100. Key screening strengths: "
            + ", ".join(strengths)
            + f". Observed LA Fitness locations: {row.LAFitnessLocations:.0f}; tracked "
            f"competitors: {row.CompetitorLocations:.0f}."
        )

    doc.add_heading("Recommended Next Actions", level=1)
    add_numbered(doc, "Validate current club rosters, closures, and planned openings in the top five counties.")
    add_numbered(doc, "Collect observed monthly dues, initiation fees, annual fees, and contract terms for a representative sample of clubs.")
    add_numbered(doc, "Replace county boundaries with 10- to 15-minute drive-time trade areas around candidate sites.")
    add_numbered(doc, "Add lease rates, daytime population, age mix, traffic counts, and competitor square footage before site underwriting.")
    add_numbered(doc, "Re-estimate the model with observed prices and validate out of sample before using it for capital allocation.")

    doc.add_heading("Limitations", level=1)
    add_bullet(doc, "Estimated target: county MembershipPrice is simulated from a disclosed formula and is not independent evidence.")
    add_bullet(doc, "Search coverage: Google Places Text Search ranks candidates and does not guarantee a complete chain inventory.")
    add_bullet(doc, "Temporal mismatch: demographics use 2020 population and 2021 income, while locations are a 2026 retrieval snapshot.")
    add_bullet(doc, "Aggregation: county averages mask neighborhood trade areas and cross-county customer travel.")
    add_bullet(doc, "Small sample: 67 counties limit model complexity and make influential observations consequential.")
    add_bullet(doc, "Association only: OLS coefficients do not identify causal effects on pricing or expansion success.")
    add_bullet(doc, "Ranking sensitivity: market-potential results depend on subjective weights and min-max scaling.")

    doc.add_heading("Reproducibility and Deliverables", level=1)
    add_bullet(doc, "fitness_market_analysis.py: API collection, integration, EDA, regression, diagnostics, predictions, and rankings.")
    add_bullet(doc, "build_reports.py: deterministic Word report generation from the analytical outputs.")
    add_bullet(doc, "data/MainDataset.csv: 67 rows and 8 required columns.")
    add_bullet(doc, "outputs/: model tables, diagnostic tables, scenario predictions, rankings, and figures.")
    add_bullet(doc, "API key files are read locally and intentionally excluded from every analytical output.")
    add_sources(doc)
    doc.save(INSIGHTS_PATH)


def main() -> None:
    data = load_inputs()
    build_eda_report(data)
    build_insights_report(data)
    print(f"Created {EDA_PATH.name}")
    print(f"Created {INSIGHTS_PATH.name}")


if __name__ == "__main__":
    main()
